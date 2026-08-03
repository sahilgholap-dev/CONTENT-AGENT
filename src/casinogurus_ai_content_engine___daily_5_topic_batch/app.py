"""FastAPI backend for the CasinoGurus content engine (production).

Replaces the stdlib ``server.py``. Same JSON API the Next.js dashboard already
speaks, now with:
  * Supabase-Auth JWT verification on every ``/api`` route (see auth.py)
  * a Postgres-backed store (see storage.py / db.py)
  * CORS for the Firebase-hosted frontend
  * a health check for the host (Render)

Run locally:
    uv run uvicorn casinogurus_ai_content_engine___daily_5_topic_batch.app:app --reload --port 8000

Endpoints (all under /api require a valid token):
    GET  /healthz
    GET  /api/formats                         -> enabled content-type/format catalog
    GET  /api/registry                        -> full catalog (admin) + task variants
    POST/PUT/DELETE /api/content-types[/{id}] -> content-type master CRUD
    POST/PUT/DELETE /api/formats[/{id}]       -> format master CRUD
    GET  /api/clients
    POST /api/clients
    GET  /api/clients/{client_id}
    PUT  /api/clients/{client_id}             -> profile edits append version N+1
    GET  /api/runs[?client_id=]
    GET  /api/batches[?client_id=]
    GET  /api/batches/{id}
    GET  /api/batches/{id}/download           -> docx ZIP
    GET  /api/latest
    GET  /api/packages/{pid}/image
    POST /api/packages/{pid}/image[?force=1]
    POST /api/packages/{pid}/feedback         -> shortlist/approve/reject event
    GET  /api/clients/{id}/learning           -> pending proposal + stats
    POST /api/clients/{id}/learning/distill   -> distil feedback into a proposal
    POST /api/clients/{id}/learning/proposals/{pid}/accept  -> new profile version
    POST /api/clients/{id}/learning/proposals/{pid}/dismiss
    POST /api/run-agent                       -> {client_id, content_type, format}
    GET  /api/agent-logs                      -> SSE
    GET/POST /api/admin/users[...]            -> portal login management
    GET  /api/portal/me|batches|batches/{id}[/download]  -> client-scoped (JWT client_id)
    GET  /api/portal/packages/{pid}/image     -> client-scoped
    POST /api/portal/packages/{pid}/feedback  -> client-scoped feedback event
"""

from __future__ import annotations

import asyncio
import io
import json
import os
import re
import subprocess
import sys
import threading
import zipfile
from datetime import datetime
from datetime import timezone as dt_timezone
from contextlib import asynccontextmanager
from typing import Literal

from fastapi import APIRouter, Depends, FastAPI, HTTPException, Query, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel, ValidationError

from casinogurus_ai_content_engine___daily_5_topic_batch.auth import (
    allowed_client_ids,
    require_admin,
    require_client,
    require_user,
    resolve_portal_scope,
)
from casinogurus_ai_content_engine___daily_5_topic_batch.db import (
    _PROJECT_ROOT,
    connection,
    init_schema,
)
from casinogurus_ai_content_engine___daily_5_topic_batch import autopilot, registry, storage
from casinogurus_ai_content_engine___daily_5_topic_batch.profile import ClientProfile
from casinogurus_ai_content_engine___daily_5_topic_batch.registry import AVAILABLE_TASK_VARIANTS
from casinogurus_ai_content_engine___daily_5_topic_batch.storage import get_image

PACKAGE = "casinogurus_ai_content_engine___daily_5_topic_batch"
LOG_PATH = os.path.join(_PROJECT_ROOT, "agent.log")

# Single-run guard: at most one crew subprocess at a time (small internal tool).
_run_state: dict = {"process": None, "log_file": None}


# --------------------------------------------------------------------------- #
# Request bodies
# --------------------------------------------------------------------------- #
class RunAgentRequest(BaseModel):
    """Defaults preserve the pre-multi-client behavior (bare POST still works)."""

    client_id: str = "casinogurus"
    content_type: str = "long_form"
    format: str = "blog"
    # Optional user-provided topic. None/empty -> the agent discovers one.
    topic: str | None = None


class ClientUpsert(BaseModel):
    id: str | None = None  # derived from display_name when omitted (POST only)
    display_name: str
    site_domain: str
    status: Literal["active", "paused", "archived"] = "active"
    profile: dict | None = None  # validated against ClientProfile when present


class FeedbackRequest(BaseModel):
    status: Literal["shortlisted", "approved", "rejected"]
    notes: str | None = None


class PortalUserCreate(BaseModel):
    email: str
    role: Literal["admin", "client"] = "client"
    client_ids: list[str] = []  # one or more clients for client logins
    client_id: str | None = None  # legacy single-client field, folded into client_ids


class PortalUserUpdate(BaseModel):
    """Edit an existing login's role / client assignments."""

    role: Literal["admin", "client"]
    client_ids: list[str] = []


class PortalRunRequest(BaseModel):
    """Portal run launch: client_id comes from the JWT, never the body."""

    content_type: str = "long_form"
    format: str = "blog"
    topic: str | None = None


class SuggestTopicsRequest(BaseModel):
    client_id: str = "casinogurus"
    content_type: str = "long_form"
    format: str = "blog"
    # Optional taste hint steering the round ("what kind of topics do you
    # have in mind"); validated like a topic (length + no brace tokens).
    hint: str | None = None


class PortalSuggestTopicsRequest(BaseModel):
    """Portal twin: client_id comes from the JWT, never the body."""

    content_type: str = "long_form"
    format: str = "blog"
    hint: str | None = None


class AutopilotConfigUpdate(BaseModel):
    """PUT payload; None fields keep their current value. Caps and timezone
    are validated server-side — never trust the client's slider."""

    paused: bool | None = None
    timezone: str | None = None
    content_types: dict | None = None


class GenerateFromSuggestionsRequest(BaseModel):
    client_id: str = "casinogurus"
    suggestion_ids: list[str]


class PortalGenerateFromSuggestionsRequest(BaseModel):
    suggestion_ids: list[str]


class ProposalAccept(BaseModel):
    # Admin may edit the distilled text before accepting; None = accept as proposed.
    text: str | None = None


class ContentTypeUpsert(BaseModel):
    id: str | None = None       # slug; derived from label on POST when omitted
    label: str
    sort_order: int = 0


class FormatUpsert(BaseModel):
    id: str | None = None       # slug; derived from label on POST when omitted
    content_type: str
    label: str
    description: str = ""
    enabled: bool = True
    task_variant: str = "default"
    pipeline: dict = {}
    stage_labels: list[str] = []
    sort_order: int = 0


def _slugify(name: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", name.lower()).strip("-")
    if not slug:
        raise HTTPException(status_code=422, detail="display_name yields an empty slug")
    return slug


def _validated_profile(profile: dict) -> dict:
    try:
        return ClientProfile.model_validate(profile).model_dump()
    except ValidationError as e:
        raise HTTPException(status_code=422, detail=f"invalid client profile: {e.errors()[:5]}")


# --------------------------------------------------------------------------- #
# Data access
# --------------------------------------------------------------------------- #
def _list_batches(client_id: str | None = None) -> list[dict]:
    where = "WHERE b.client_id = %(client_id)s" if client_id else ""
    with connection() as conn:
        rows = conn.execute(
            f"""SELECT b.id, b.batch_date, b.total_packages, b.ready_for_review_count,
                      b.needs_review_count, b.source, b.ingested_at,
                      b.client_id, b.content_type, b.format,
                      c.display_name AS client_name,
                      COUNT(p.package_id) AS package_count
               FROM batches b
               LEFT JOIN packages p ON p.batch_id = b.id
               LEFT JOIN clients c ON c.id = b.client_id
               {where}
               GROUP BY b.id, c.display_name
               ORDER BY b.ingested_at DESC, b.id DESC""",
            {"client_id": client_id},
        ).fetchall()
        return [dict(r) for r in rows]


def _merge_feedback(batch: dict) -> dict:
    """Attach the latest reviewer feedback event to each package dict so the
    viewer needs no second request. Best-effort: never fails the batch load."""
    try:
        packages = batch.get("packages") or []
        ids = [p.get("package_id") for p in packages if p.get("package_id")]
        reviews = storage.latest_reviews_for_packages(ids)
        for pkg in packages:
            fb = reviews.get(pkg.get("package_id"))
            if fb:
                pkg["feedback"] = {k: v for k, v in fb.items() if k != "package_id"}
    except Exception:
        pass
    return batch


def _get_batch(batch_id: int) -> dict | None:
    with connection() as conn:
        row = conn.execute(
            "SELECT id, raw_json FROM batches WHERE id = %s", (batch_id,)
        ).fetchone()
    if not row:
        return None
    # raw_json is JSONB, so psycopg returns a dict already (no json.loads).
    batch = row["raw_json"]
    if isinstance(batch, dict):
        batch["id"] = row["id"]
        batch = _merge_feedback(batch)
    return batch


def _latest_batch() -> dict | None:
    with connection() as conn:
        row = conn.execute(
            "SELECT id, raw_json FROM batches ORDER BY ingested_at DESC, id DESC LIMIT 1"
        ).fetchone()
    if not row:
        return None
    batch = row["raw_json"]
    if isinstance(batch, dict):
        batch["id"] = row["id"]
    return batch


def _image_payload(row: dict | None) -> dict | None:
    if not row:
        return None
    return {
        "package_id": row.get("package_id"),
        "status": row.get("status"),
        "image_b64": row.get("image_b64"),
        "mime_type": row.get("mime_type"),
        "alt_text": row.get("alt_text"),
        "prompt": row.get("prompt"),
        "model": row.get("model"),
        "size": row.get("size"),
        "error": row.get("error"),
        "created_at": row.get("created_at"),
    }


# --------------------------------------------------------------------------- #
# DOCX ZIP export (ported verbatim from the old server._send_zip)
# --------------------------------------------------------------------------- #
def build_batch_zip(batch: dict) -> bytes:
    from bs4 import BeautifulSoup
    import docx

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        for idx, pkg in enumerate(batch.get("packages", [])):
            doc = docx.Document()
            title = pkg.get("topic") or pkg.get("primary_keyword") or f"Topic_{idx+1}"
            doc.add_heading(title, 0)

            doc.add_heading("Metadata", level=1)
            doc.add_paragraph(f"Primary Keyword: {pkg.get('primary_keyword', '')}")

            draft = pkg.get("draft", {})
            doc.add_paragraph(f"Meta Description: {draft.get('meta_description', '')}")

            doc.add_heading("Content", level=1)
            html_content = draft.get("body_html", "")
            if html_content:
                soup = BeautifulSoup(html_content, "html.parser")
                for element in soup.find_all(["h2", "h3", "p", "ul", "ol", "li"]):
                    if element.name == "h2":
                        doc.add_heading(element.get_text(), level=2)
                    elif element.name == "h3":
                        doc.add_heading(element.get_text(), level=3)
                    elif element.name == "p":
                        doc.add_paragraph(element.get_text())
                    elif element.name in ["ul", "ol"]:
                        pass
                    elif element.name == "li":
                        doc.add_paragraph(element.get_text(), style="List Bullet")

            flags = draft.get("verification_flags", [])
            if flags:
                doc.add_heading("Verification Flags", level=1)
                for f in flags:
                    flag_str = f.get("flag") if isinstance(f, dict) else str(f)
                    doc.add_paragraph(flag_str, style="List Bullet")

            compliance = pkg.get("compliance_scorecard", {})
            if compliance:
                doc.add_heading("Compliance Checks", level=1)
                doc.add_paragraph(f"Overall Verdict: {compliance.get('overall_verdict', 'Unknown')}")

                blocking = compliance.get("blocking_failures", [])
                if blocking:
                    doc.add_heading("Blocking Failures", level=2)
                    for b in blocking:
                        if isinstance(b, str):
                            doc.add_paragraph(b, style="List Bullet")
                        else:
                            name = b.get("check_name") or b.get("item") or b.get("check") or b.get("name") or "Unknown Check"
                            sev = b.get("severity")
                            sev_str = f" ({sev})" if sev else ""
                            doc.add_heading(f"{name}{sev_str}", level=3)
                            if b.get("violation"):
                                doc.add_paragraph(f"Violation: {b.get('violation')}")
                            if b.get("remediation"):
                                doc.add_paragraph(f"Remediation: {b.get('remediation')}")

                checks = compliance.get("checks", [])
                if checks:
                    doc.add_heading("All Checks", level=2)
                    for c in checks:
                        name = c.get("check_name") or c.get("item") or c.get("check") or c.get("name") or "Unknown Check"
                        verdict = c.get("verdict") or c.get("result") or ""
                        sev = c.get("severity")
                        sev_str = f" [{sev}]" if sev else ""
                        doc.add_paragraph(f"{name}{sev_str}: {verdict}", style="List Bullet")
                        details = c.get("offending_text") or c.get("offendingText") or c.get("violation") or c.get("details") or ""
                        if details:
                            doc.add_paragraph(f"   Details: {details}")

            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)

            filename = "".join([c for c in title if c.isalpha() or c.isdigit() or c == " "]).rstrip()
            filename = filename.replace(" ", "_") + ".docx"
            zip_file.writestr(filename, doc_buffer.getvalue())

    zip_buffer.seek(0)
    return zip_buffer.read()


# --------------------------------------------------------------------------- #
# Routes
# --------------------------------------------------------------------------- #
# Everything on this router is internal-team functionality: it requires
# app_metadata.role == 'admin'. Client-portal logins use the /api/portal
# router further below, which scopes every query to the token's client_id.
api = APIRouter(prefix="/api", dependencies=[Depends(require_admin)])


@api.get("/formats")
def list_formats():
    """Enabled content-type/format catalog (cascading), for the run modal."""
    return jsonable(storage.serialisable_registry(enabled_only=True))


# --- Content-type / format master (admin CRUD) ---------------------------- #
@api.get("/registry")
def get_registry():
    """Full catalog for the management UI: all content types + all formats
    (including disabled), plus the task variants the pipeline code supports."""
    return jsonable(
        {
            "content_types": storage.list_content_types(),
            "formats": storage.list_formats(enabled_only=False),
            "task_variants": AVAILABLE_TASK_VARIANTS,
        }
    )


@api.post("/content-types", status_code=201)
def create_content_type(body: ContentTypeUpsert):
    ct_id = body.id or _slugify(body.label)
    if any(c["id"] == ct_id for c in storage.list_content_types()):
        raise HTTPException(status_code=409, detail=f"content type '{ct_id}' already exists")
    return jsonable(storage.upsert_content_type(ct_id, body.label, body.sort_order))


@api.put("/content-types/{ct_id}")
def update_content_type(ct_id: str, body: ContentTypeUpsert):
    if not any(c["id"] == ct_id for c in storage.list_content_types()):
        raise HTTPException(status_code=404, detail=f"content type '{ct_id}' not found")
    return jsonable(storage.upsert_content_type(ct_id, body.label, body.sort_order))


@api.delete("/content-types/{ct_id}", status_code=204)
def remove_content_type(ct_id: str):
    storage.delete_content_type(ct_id)  # cascades to its formats
    return Response(status_code=204)


@api.post("/formats", status_code=201)
def create_format(body: FormatUpsert):
    fmt_id = body.id or _slugify(body.label)
    if storage.get_format_row(fmt_id):
        raise HTTPException(status_code=409, detail=f"format '{fmt_id}' already exists")
    _validate_format_body(body)
    return jsonable(
        storage.upsert_format(
            fmt_id, body.content_type, body.label, body.description, body.enabled,
            body.task_variant, body.pipeline, body.stage_labels, body.sort_order,
        )
    )


@api.put("/formats/{fmt_id}")
def update_format(fmt_id: str, body: FormatUpsert):
    if not storage.get_format_row(fmt_id):
        raise HTTPException(status_code=404, detail=f"format '{fmt_id}' not found")
    _validate_format_body(body)
    return jsonable(
        storage.upsert_format(
            fmt_id, body.content_type, body.label, body.description, body.enabled,
            body.task_variant, body.pipeline, body.stage_labels, body.sort_order,
        )
    )


@api.delete("/formats/{fmt_id}", status_code=204)
def remove_format(fmt_id: str):
    storage.delete_format(fmt_id)
    return Response(status_code=204)


def _validate_format_body(body: FormatUpsert) -> None:
    if not any(c["id"] == body.content_type for c in storage.list_content_types()):
        raise HTTPException(status_code=422, detail=f"unknown content_type '{body.content_type}'")
    if body.task_variant not in AVAILABLE_TASK_VARIANTS:
        raise HTTPException(
            status_code=422,
            detail=f"task_variant '{body.task_variant}' is not implemented in the pipeline "
                   f"(available: {AVAILABLE_TASK_VARIANTS})",
        )


@api.get("/clients")
def list_clients():
    return jsonable(storage.list_clients())


@api.post("/clients", status_code=201)
def create_client(body: ClientUpsert):
    client_id = body.id or _slugify(body.display_name)
    if storage.get_client(client_id):
        raise HTTPException(status_code=409, detail=f"client '{client_id}' already exists")
    if body.profile is None:
        raise HTTPException(status_code=422, detail="profile is required when creating a client")
    profile = _validated_profile(body.profile)
    storage.upsert_client(client_id, body.display_name, body.site_domain, body.status)
    storage.insert_profile_version(client_id, profile, created_by="api")
    return jsonable(storage.get_client(client_id))


@api.get("/clients/{client_id}")
def get_client(client_id: str):
    client = storage.get_client(client_id)
    if not client:
        raise HTTPException(status_code=404, detail=f"client '{client_id}' not found")
    return jsonable(client)


@api.put("/clients/{client_id}")
def update_client(client_id: str, body: ClientUpsert):
    existing = storage.get_client(client_id)
    if not existing:
        raise HTTPException(status_code=404, detail=f"client '{client_id}' not found")
    storage.upsert_client(client_id, body.display_name, body.site_domain, body.status)
    if body.profile is not None:
        # Profiles are append-only: every edit becomes version N+1; in-flight
        # runs keep the version they pinned at kickoff.
        storage.insert_profile_version(client_id, _validated_profile(body.profile), created_by="api")
    return jsonable(storage.get_client(client_id))


@api.get("/runs")
def list_runs(client_id: str | None = Query(default=None)):
    return jsonable(storage.list_runs(client_id=client_id))


@api.post("/packages/{pid}/feedback")
def package_feedback(pid: str, body: FeedbackRequest, user: dict = Depends(require_user)):
    reviewer = (user or {}).get("email") or (user or {}).get("sub")
    try:
        row = storage.add_package_review(pid, body.status, body.notes, reviewer)
    except KeyError:
        raise HTTPException(status_code=404, detail=f"package '{pid}' not found")
    return jsonable(row)


# --------------------------------------------------------------------------- #
# Portal user management (Supabase Auth Admin API; admin-only like the rest
# of this router). Client logins carry app_metadata {role, client_id}.
# --------------------------------------------------------------------------- #
def _sb():
    from casinogurus_ai_content_engine___daily_5_topic_batch import supabase_admin

    return supabase_admin


def _user_row(u: dict) -> dict:
    meta = u.get("app_metadata") or {}
    return {
        "id": u.get("id"),
        "email": u.get("email"),
        "role": meta.get("role"),
        "client_ids": allowed_client_ids(meta),
        "created_at": u.get("created_at"),
        "last_sign_in_at": u.get("last_sign_in_at"),
        "disabled": bool(u.get("banned_until")),
    }


def _sb_call(fn, *args, **kwargs):
    from casinogurus_ai_content_engine___daily_5_topic_batch.supabase_admin import (
        SupabaseAdminError,
    )

    try:
        return fn(*args, **kwargs)
    except SupabaseAdminError as e:
        raise HTTPException(status_code=e.status_code if e.status_code < 500 else 502, detail=str(e))


@api.get("/admin/users")
def list_portal_users():
    return [_user_row(u) for u in _sb_call(_sb().list_users)]


def _validated_client_ids(role: str, client_ids: list[str]) -> list[str] | None:
    """Normalised, validated client list for a login: None for admins, a
    deduped non-empty list of EXISTING client ids for client logins."""
    if role != "client":
        return None
    seen: set[str] = set()
    ids: list[str] = []
    for raw in client_ids:
        cid = (raw or "").strip()
        if cid and cid not in seen:
            seen.add(cid)
            ids.append(cid)
    if not ids:
        raise HTTPException(
            status_code=422, detail="at least one client_id is required for client logins"
        )
    for cid in ids:
        if not storage.get_client(cid):
            raise HTTPException(status_code=404, detail=f"client '{cid}' not found")
    return ids


@api.post("/admin/users")
def create_portal_user(body: PortalUserCreate):
    email = body.email.strip().lower()
    requested = list(body.client_ids)
    if body.client_id:  # legacy single-client callers
        requested.append(body.client_id)
    client_ids = _validated_client_ids(body.role, requested)
    sb = _sb()
    temp_password = sb.generate_temp_password()
    user = _sb_call(sb.create_user, email, temp_password, body.role, client_ids)
    # temp_password is returned ONCE, here, and never stored.
    return {"user": _user_row(user), "temp_password": temp_password}


@api.put("/admin/users/{user_id}")
def update_portal_user(user_id: str, body: PortalUserUpdate, user: dict = Depends(require_user)):
    if user_id == user.get("sub"):
        raise HTTPException(status_code=409, detail="You cannot edit your own account.")
    client_ids = _validated_client_ids(body.role, body.client_ids)
    return _user_row(_sb_call(_sb().set_role, user_id, body.role, client_ids))


@api.post("/admin/users/{user_id}/reset-password")
def reset_portal_user_password(user_id: str):
    sb = _sb()
    temp_password = sb.generate_temp_password()
    _sb_call(sb.set_password, user_id, temp_password)
    return {"temp_password": temp_password}


@api.post("/admin/users/{user_id}/disable")
def disable_portal_user(user_id: str, user: dict = Depends(require_user)):
    if user_id == user.get("sub"):
        raise HTTPException(status_code=409, detail="You cannot disable your own account.")
    return _user_row(_sb_call(_sb().set_banned, user_id, True))


@api.post("/admin/users/{user_id}/enable")
def enable_portal_user(user_id: str):
    return _user_row(_sb_call(_sb().set_banned, user_id, False))


# --------------------------------------------------------------------------- #
# Learning loop (human-gated learned_style distillation)
# --------------------------------------------------------------------------- #
@api.get("/clients/{client_id}/learning")
def learning_state(client_id: str):
    """Pending proposal, unprocessed-event count, and approval-rate stats."""
    if not storage.get_client(client_id):
        raise HTTPException(status_code=404, detail=f"client '{client_id}' not found")
    state = storage.get_learning_state(client_id)
    state["stats"] = storage.approval_stats_by_profile_version(client_id)
    return jsonable(state)


@api.post("/clients/{client_id}/learning/distill")
def learning_distill(client_id: str):
    """Analyse new review events and park a pending learned_style proposal."""
    if not storage.get_client(client_id):
        raise HTTPException(status_code=404, detail=f"client '{client_id}' not found")
    from casinogurus_ai_content_engine___daily_5_topic_batch import learning

    try:
        result = learning.distill_client(client_id)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"distillation failed: {e}")
    return jsonable(result)


@api.post("/clients/{client_id}/learning/proposals/{proposal_id}/accept")
def learning_accept(
    client_id: str, proposal_id: int, body: ProposalAccept, user: dict = Depends(require_user)
):
    """Accept a proposal (optionally with admin edits) -> new profile version."""
    decided_by = (user or {}).get("email") or (user or {}).get("sub")
    row = storage.decide_learning_proposal(proposal_id, client_id, "accepted", decided_by)
    if not row:
        raise HTTPException(status_code=404, detail="no matching pending proposal")
    client = storage.get_client(client_id)
    if not client or not client.get("profile"):
        raise HTTPException(status_code=409, detail=f"client '{client_id}' has no profile to update")
    profile = dict(client["profile"])
    profile["learned_style"] = (body.text if body.text is not None else row["proposed_text"]).strip()
    version = storage.insert_profile_version(
        client_id, _validated_profile(profile), created_by=f"learning-loop:{decided_by}"
    )
    return jsonable({"proposal": row, "profile_version": version})


@api.post("/clients/{client_id}/learning/proposals/{proposal_id}/dismiss")
def learning_dismiss(client_id: str, proposal_id: int, user: dict = Depends(require_user)):
    decided_by = (user or {}).get("email") or (user or {}).get("sub")
    row = storage.decide_learning_proposal(proposal_id, client_id, "dismissed", decided_by)
    if not row:
        raise HTTPException(status_code=404, detail="no matching pending proposal")
    return jsonable({"proposal": row})


@api.get("/batches")
def list_batches(client_id: str | None = Query(default=None)):
    return jsonable(_list_batches(client_id))


@api.get("/latest")
def latest_batch():
    batch = _latest_batch()
    if batch is None:
        raise HTTPException(status_code=404, detail="no batches stored")
    return batch


@api.get("/batches/{batch_id}")
def get_batch(batch_id: int):
    batch = _get_batch(batch_id)
    if batch is None:
        raise HTTPException(status_code=404, detail=f"batch {batch_id} not found")
    return batch


@api.get("/batches/{batch_id}/download")
def download_batch(batch_id: int):
    batch = _get_batch(batch_id)
    if batch is None:
        raise HTTPException(status_code=404, detail=f"batch {batch_id} not found")
    body = build_batch_zip(batch)
    batch_date = batch.get("batch_date", "download")
    return Response(
        content=body,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="batch_{batch_date}.zip"'},
    )


@api.get("/packages/{pid}/image")
def get_package_image(pid: str):
    payload = _image_payload(get_image(pid))
    if payload is None or payload.get("status") != "ok" or not payload.get("image_b64"):
        body = payload or {"package_id": pid, "status": "none"}
        return JSONResponse(status_code=404, content=jsonable(body))
    return payload


@api.post("/packages/{pid}/image")
def generate_package_image(pid: str, force: bool = Query(default=False)):
    from casinogurus_ai_content_engine___daily_5_topic_batch.images import generate_for_package

    try:
        row = generate_for_package(pid, force=force)
    except ValueError as e:  # unknown package id
        raise HTTPException(status_code=404, detail=str(e))
    payload = _image_payload(row)
    status_code = 200 if payload and payload.get("status") == "ok" else 502
    return JSONResponse(status_code=status_code, content=jsonable(payload or {"package_id": pid, "status": "error"}))


# --------------------------------------------------------------------------- #
# Client portal API. Every endpoint is scoped to the client_id baked into the
# caller's JWT (app_metadata.client_id) -- NEVER to a query param from a
# client login. Admins may also call these for support/testing by passing
# ?client_id= explicitly. No run-agent, registry, or profile access here.
# --------------------------------------------------------------------------- #
portal = APIRouter(prefix="/api/portal", dependencies=[Depends(require_client)])


def _portal_cid(user: dict, client_id: str | None = None) -> str:
    """The effective client scope: validated against the JWT's client list
    for client logins, or the explicit ?client_id= param for admins."""
    return resolve_portal_scope(user.get("portal_client_ids"), client_id)


def _batch_client_id(batch_id: int) -> str | None:
    with connection() as conn:
        row = conn.execute("SELECT client_id FROM batches WHERE id = %s", (batch_id,)).fetchone()
        return row["client_id"] if row else None


def _package_client_id(pid: str) -> str | None:
    with connection() as conn:
        row = conn.execute(
            "SELECT client_id FROM packages WHERE package_id = %s", (pid,)
        ).fetchone()
        return row["client_id"] if row else None


def _own_batch_or_404(batch_id: int, cid: str) -> dict:
    # Same 404 whether the batch doesn't exist or belongs to another client --
    # portal callers can't probe other clients' batch ids.
    if _batch_client_id(batch_id) != cid:
        raise HTTPException(status_code=404, detail=f"batch {batch_id} not found")
    batch = _get_batch(batch_id)
    if batch is None:
        raise HTTPException(status_code=404, detail=f"batch {batch_id} not found")
    return batch


def _own_package_or_404(pid: str, cid: str) -> None:
    if _package_client_id(pid) != cid:
        raise HTTPException(status_code=404, detail=f"package '{pid}' not found")


@portal.get("/me")
def portal_me(user: dict = Depends(require_client), client_id: str | None = Query(default=None)):
    allowed = user.get("portal_client_ids")
    if allowed is None:
        # Admin browsing a specific client's portal view.
        allowed = [_portal_cid(user, client_id)]
    clients = []
    for cid in allowed:
        c = storage.get_client(cid)
        if c:
            clients.append(
                {
                    "id": c["id"],
                    "display_name": c["display_name"],
                    "site_domain": c["site_domain"],
                    "status": c["status"],
                }
            )
    if not clients:
        raise HTTPException(status_code=404, detail="client not found")
    # Deliberately NOT the profile: it contains internal prompt engineering.
    return jsonable({"email": user.get("email"), "clients": clients})


@portal.get("/batches")
def portal_batches(user: dict = Depends(require_client), client_id: str | None = Query(default=None)):
    return jsonable(_list_batches(_portal_cid(user, client_id)))


@portal.get("/batches/{batch_id}")
def portal_batch(
    batch_id: int, user: dict = Depends(require_client), client_id: str | None = Query(default=None)
):
    return _own_batch_or_404(batch_id, _portal_cid(user, client_id))


@portal.get("/batches/{batch_id}/download")
def portal_download(
    batch_id: int, user: dict = Depends(require_client), client_id: str | None = Query(default=None)
):
    batch = _own_batch_or_404(batch_id, _portal_cid(user, client_id))
    body = build_batch_zip(batch)
    batch_date = batch.get("batch_date", "download")
    return Response(
        content=body,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="batch_{batch_date}.zip"'},
    )


@portal.get("/packages/{pid}/image")
def portal_package_image(
    pid: str, user: dict = Depends(require_client), client_id: str | None = Query(default=None)
):
    _own_package_or_404(pid, _portal_cid(user, client_id))
    payload = _image_payload(get_image(pid))
    if payload is None or payload.get("status") != "ok" or not payload.get("image_b64"):
        body = payload or {"package_id": pid, "status": "none"}
        return JSONResponse(status_code=404, content=jsonable(body))
    return payload


@portal.post("/packages/{pid}/feedback")
def portal_feedback(
    pid: str,
    body: FeedbackRequest,
    user: dict = Depends(require_client),
    client_id: str | None = Query(default=None),
):
    """Client feedback -- the same event stream the learning loop distils."""
    _own_package_or_404(pid, _portal_cid(user, client_id))
    reviewer = user.get("email") or user.get("sub")
    try:
        row = storage.add_package_review(pid, body.status, body.notes, reviewer)
    except KeyError:
        raise HTTPException(status_code=404, detail=f"package '{pid}' not found")
    return jsonable(row)


@portal.get("/pieces")
def portal_pieces(user: dict = Depends(require_client), client_id: str | None = Query(default=None)):
    """Flat piece list (one per package, newest first) for the Content
    Studio Drafts/Approved views. state: drafted|shortlisted|approved|rejected."""
    return jsonable(storage.list_client_pieces(_portal_cid(user, client_id)))


@portal.get("/pieces/{pid}")
def portal_piece(
    pid: str, user: dict = Depends(require_client), client_id: str | None = Query(default=None)
):
    piece = storage.get_client_piece(_portal_cid(user, client_id), pid)
    if piece is None:
        raise HTTPException(status_code=404, detail=f"piece '{pid}' not found")
    return jsonable(piece)


@portal.get("/formats")
def portal_formats():
    """Enabled content-type/format catalog for the portal run modal (same
    non-sensitive payload the admin run modal uses)."""
    return jsonable(storage.serialisable_registry(enabled_only=True))


@portal.post("/run-agent")
def portal_run_agent(
    body: PortalRunRequest,
    user: dict = Depends(require_client),
    client_id: str | None = Query(default=None),
):
    """Clients can generate content for their own business; the run is always
    scoped to the JWT's client_id."""
    cid = _portal_cid(user, client_id)
    return _start_agent_run(cid, body.content_type, body.format, body.topic)


@portal.post("/suggest-topics")
def portal_suggest_topics(
    body: PortalSuggestTopicsRequest,
    user: dict = Depends(require_client),
    client_id: str | None = Query(default=None),
):
    cid = _portal_cid(user, client_id)
    return _start_agent_run(cid, body.content_type, body.format, body.hint, kind="suggest")


@portal.get("/topic-suggestions")
def portal_topic_suggestions(
    format: str = Query(...),
    user: dict = Depends(require_client),
    client_id: str | None = Query(default=None),
):
    return jsonable(storage.list_topic_suggestions(_portal_cid(user, client_id), format))


@portal.post("/generate-from-suggestions")
def portal_generate_from_suggestions(
    body: PortalGenerateFromSuggestionsRequest,
    user: dict = Depends(require_client),
    client_id: str | None = Query(default=None),
):
    return _generate_from_suggestions(_portal_cid(user, client_id), body.suggestion_ids)


@portal.get("/autopilot/config")
def portal_autopilot_config(
    user: dict = Depends(require_client), client_id: str | None = Query(default=None)
):
    return jsonable(_autopilot_config_payload(_portal_cid(user, client_id)))


@portal.put("/autopilot/config")
def portal_autopilot_config_update(
    body: AutopilotConfigUpdate,
    user: dict = Depends(require_client),
    client_id: str | None = Query(default=None),
):
    return jsonable(_update_autopilot_config(_portal_cid(user, client_id), body))


@portal.post("/autopilot/pause")
def portal_autopilot_pause(
    user: dict = Depends(require_client), client_id: str | None = Query(default=None)
):
    cid = _portal_cid(user, client_id)
    return jsonable(_update_autopilot_config(cid, AutopilotConfigUpdate(paused=True)))


@portal.post("/autopilot/resume")
def portal_autopilot_resume(
    user: dict = Depends(require_client), client_id: str | None = Query(default=None)
):
    cid = _portal_cid(user, client_id)
    return jsonable(_update_autopilot_config(cid, AutopilotConfigUpdate(paused=False)))


@portal.get("/autopilot/queue")
def portal_autopilot_queue(
    user: dict = Depends(require_client), client_id: str | None = Query(default=None)
):
    return jsonable(storage.list_autopilot_queue(_portal_cid(user, client_id)))


@portal.post("/autopilot/queue/{queue_id}/approve")
def portal_autopilot_queue_approve(
    queue_id: str, user: dict = Depends(require_client), client_id: str | None = Query(default=None)
):
    return jsonable(_queue_approve(_portal_cid(user, client_id), queue_id))


@portal.post("/autopilot/queue/{queue_id}/swap")
def portal_autopilot_queue_swap(
    queue_id: str, user: dict = Depends(require_client), client_id: str | None = Query(default=None)
):
    return jsonable(_queue_swap(_portal_cid(user, client_id), queue_id))


@portal.post("/autopilot/queue/{queue_id}/skip")
def portal_autopilot_queue_skip(
    queue_id: str, user: dict = Depends(require_client), client_id: str | None = Query(default=None)
):
    return jsonable(_queue_skip(_portal_cid(user, client_id), queue_id))


@portal.get("/run-progress")
def portal_run_progress(
    user: dict = Depends(require_client), client_id: str | None = Query(default=None)
):
    """Stage progress of the CURRENT agent run, only if it belongs to the
    caller's client. Parsed server-side from the run log (the [AGENT_RUN]
    header + '[AGENT_PROGRESS] Task Completed' markers the crew emits) so
    clients never see the raw log stream — another client's run returns the
    same empty shape as no run at all."""
    empty = {"active": False, "stage": 0, "total": 0, "label": None}
    cid = _portal_cid(user, client_id)
    proc = _run_state["process"]
    running = proc is not None and proc.poll() is None
    if not os.path.exists(LOG_PATH):
        return empty
    header: dict | None = None
    completed = 0
    try:
        with open(LOG_PATH, "r", encoding="utf-8", errors="replace") as f:
            for line in f:
                if header is None and line.startswith("[AGENT_RUN] "):
                    try:
                        header = json.loads(line[len("[AGENT_RUN] "):])
                    except Exception:
                        header = {}
                if "[AGENT_PROGRESS] Task Completed" in line:
                    completed += 1
    except OSError:
        return empty
    if not header or header.get("client_id") != cid:
        return empty
    labels = [str(x) for x in (header.get("stage_labels") or [])]
    total = len(labels) or 5
    stage = min(completed, total)
    return {
        "active": running,
        "stage": stage,
        "total": total,
        "label": labels[stage] if (running and stage < len(labels)) else None,
        "run_id": header.get("run_id"),
    }


@portal.get("/runs")
def portal_runs(user: dict = Depends(require_client), client_id: str | None = Query(default=None)):
    """The caller's run history, trimmed to non-internal fields (the portal
    polls this to show progress and refresh batches when a run finishes)."""
    cid = _portal_cid(user, client_id)
    rows = storage.list_runs(client_id=cid)
    keep = ("id", "status", "content_type", "format", "topic",
            "created_at", "started_at", "finished_at", "batch_id")
    return jsonable([{k: r.get(k) for k in keep} for r in rows])


def _tee_output(process, log_file):
    try:
        for line in iter(process.stdout.readline, b""):
            sys.stdout.buffer.write(line)
            sys.stdout.flush()
            log_file.write(line.decode("utf-8", errors="replace"))
            log_file.flush()
    except Exception as e:
        print(f"Error in tee thread: {e}")
    finally:
        process.stdout.close()
        log_file.close()


def _validate_suggestion_selection(
    rows: list[dict], suggestion_ids: list[str], client_id: str, cap: int
) -> tuple[str, str, list[str]]:
    """Validate a shortlist selection; returns (content_type, format,
    ordered_topics) in the user's tick order. 422 on any violation, with a
    message the modal can show verbatim."""
    if not suggestion_ids:
        raise HTTPException(status_code=422, detail="select at least one suggested topic")
    if len(set(suggestion_ids)) != len(suggestion_ids):
        raise HTTPException(status_code=422, detail="duplicate suggestion ids in selection")
    by_id = {str(r["id"]): r for r in rows}
    missing = [i for i in suggestion_ids if i not in by_id]
    if missing:
        raise HTTPException(status_code=422, detail=f"unknown suggestion id(s): {missing}")
    ordered = [by_id[i] for i in suggestion_ids]
    if any(r["client_id"] != client_id for r in ordered):
        raise HTTPException(status_code=422, detail="suggestion belongs to a different client")
    if len({r["format"] for r in ordered}) > 1:
        raise HTTPException(status_code=422, detail="all selected topics must share one format")
    used = [str(r["id"]) for r in ordered if r["status"] != "suggested"]
    if used:
        raise HTTPException(
            status_code=422, detail="some selected topics were already used — refresh the list"
        )
    if len(ordered) > cap:
        raise HTTPException(
            status_code=422, detail=f"this format allows at most {cap} topics per run"
        )
    return ordered[0]["content_type"], ordered[0]["format"], [r["topic"] for r in ordered]


def _generate_from_suggestions(client_id: str, suggestion_ids: list[str]) -> dict:
    """Shared launch path for both generate-from-suggestions surfaces."""
    rows = storage.get_topic_suggestions_by_ids(suggestion_ids)
    fmt = rows[0]["format"] if rows else ""
    spec = storage.resolve_format_spec(fmt) if fmt else None
    cap = registry.max_per_run(spec) if spec else 1
    content_type, format_id, topics = _validate_suggestion_selection(
        rows, suggestion_ids, client_id, cap
    )
    result = _start_agent_run(client_id, content_type, format_id, None, kind="generate", topics=topics)
    storage.set_suggestions_status(suggestion_ids, "selected", generate_run_id=result["run_id"])
    return result


def _start_agent_run(
    client_id: str,
    content_type: str,
    format_id: str,
    raw_topic: str | None,
    kind: str = "generate",
    topics: list[str] | None = None,
    origin: str = "manual",
) -> dict:
    """Shared run-launch path for the admin endpoint and the client portal.

    All validation lives here so both surfaces behave identically; the only
    difference is where client_id comes from (request body vs JWT).
    kind='suggest' launches a discovery-only suggestion round (raw_topic is
    the taste hint); topics pins a shortlist onto a generate run."""
    proc = _run_state["process"]
    if proc is not None and proc.poll() is None:
        raise HTTPException(status_code=409, detail="Agent is already running")

    # Validate the requested format against the DB catalog.
    spec = storage.resolve_format_spec(format_id)
    if spec is None:
        raise HTTPException(status_code=422, detail=f"unknown format '{format_id}' (see /api/formats)")
    if not spec.enabled:
        raise HTTPException(status_code=422, detail=f"format '{format_id}' is not enabled")
    if spec.content_type != content_type:
        raise HTTPException(
            status_code=422,
            detail=f"format '{format_id}' belongs to content type '{spec.content_type}', not '{content_type}'",
        )

    # Validate the client and pin its current profile version via a runs row.
    client = storage.get_client(client_id)
    if not client:
        raise HTTPException(status_code=404, detail=f"client '{client_id}' not found")
    if client["status"] != "active":
        raise HTTPException(status_code=409, detail=f"client '{client_id}' is {client['status']}")
    if not client["profile_version"]:
        raise HTTPException(status_code=409, detail=f"client '{client_id}' has no profile yet")

    # Optional user-provided topic: trimmed, bounded, and interpolation-safe
    # (a {token} in the topic would be substituted into the prompts or crash
    # kickoff — same rule as profile text).
    topic = (raw_topic or "").strip() or None
    if topic:
        if len(topic) > 300:
            raise HTTPException(status_code=422, detail="topic must be at most 300 characters")
        if re.search(r"\{[A-Za-z_][A-Za-z0-9_\-]*\}", topic):
            raise HTTPException(status_code=422, detail="topic must not contain {placeholder}-style tokens")

    run_row = storage.create_run(client_id, content_type, format_id, topic=topic, kind=kind, topics=topics, origin=origin)

    # With a pinned topic (user-provided or shortlist) the first pipeline task
    # STRUCTURES the given topic instead of discovering one — label the
    # terminal stage honestly so it doesn't look like discovery re-ran.
    stage_labels = ["Topic Suggestions"] if kind == "suggest" else list(spec.stage_labels)
    if kind != "suggest" and (topic or topics) and stage_labels and stage_labels[0] == "Topic Discovery":
        stage_labels[0] = "Structuring Selected Topic" + ("s" if topics and len(topics) > 1 else "")

    log_file = open(LOG_PATH, "w", encoding="utf-8")
    # First log line self-describes the run so the SSE terminal can label the
    # stages and header without another request (SSE replays from file start).
    log_file.write(
        "[AGENT_RUN] "
        + json.dumps(
            {
                "run_id": str(run_row["id"]),
                "client_id": client["id"],
                "client_name": client["display_name"],
                "content_type": spec.content_type,
                "format": spec.id,
                "topic": topic,
                "kind": kind,
                "stage_labels": stage_labels,
            }
        )
        + "\n"
    )
    log_file.flush()

    cmd = [sys.executable, "-u", "-m", f"{PACKAGE}.main", "run", "--run-id", str(run_row["id"])]
    env = os.environ.copy()
    env["PYTHONIOENCODING"] = "utf-8"
    env["PYTHONUNBUFFERED"] = "1"
    try:
        process = subprocess.Popen(
            cmd, cwd=_PROJECT_ROOT, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, env=env
        )
        threading.Thread(target=_tee_output, args=(process, log_file), daemon=True).start()
    except Exception as e:
        log_file.close()
        storage.update_run(run_row["id"], status="failed", error=f"spawn failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    _run_state.update(process=process, log_file=log_file)
    return {
        "status": "started",
        "message": "Agent execution started in background.",
        "run_id": str(run_row["id"]),
        "client_id": client["id"],
        "format": spec.id,
    }


@api.post("/run-agent")
def run_agent(body: RunAgentRequest | None = None):
    body = body or RunAgentRequest()
    return _start_agent_run(body.client_id, body.content_type, body.format, body.topic)


@api.post("/suggest-topics")
def suggest_topics(body: SuggestTopicsRequest):
    """Start a discovery-only suggestion round (run kind: suggest)."""
    return _start_agent_run(body.client_id, body.content_type, body.format, body.hint, kind="suggest")


@api.get("/topic-suggestions")
def topic_suggestions(client_id: str = Query(...), format: str = Query(...)):
    """Available (status=suggested) topics for the client+format, newest first."""
    return jsonable(storage.list_topic_suggestions(client_id, format))


@api.post("/generate-from-suggestions")
def generate_from_suggestions(body: GenerateFromSuggestionsRequest):
    return _generate_from_suggestions(body.client_id, body.suggestion_ids)


# --- Autopilot (shared handlers; admin routes here, portal twins below) ---- #

def _autopilot_config_payload(client_id: str) -> dict:
    cfg = storage.get_autopilot_config(client_id) or {}
    return {
        "paused": bool(cfg.get("paused", False)),
        "timezone": cfg.get("timezone") or autopilot.DEFAULT_TIMEZONE,
        "content_types": autopilot.normalize_config(cfg.get("content_types")),
        "caps": autopilot.AUTOPILOT_CAPS,
    }


def _update_autopilot_config(client_id: str, body: AutopilotConfigUpdate) -> dict:
    if not storage.get_client(client_id):
        raise HTTPException(status_code=404, detail=f"client '{client_id}' not found")
    if body.timezone is not None and not autopilot.validate_timezone(body.timezone):
        raise HTTPException(status_code=400, detail=f"unknown timezone '{body.timezone}'")
    normalized = None
    if body.content_types is not None:
        problems = autopilot.validate_config(body.content_types)
        if problems:
            raise HTTPException(status_code=400, detail="; ".join(problems))
        normalized = autopilot.normalize_config(body.content_types)
    storage.upsert_autopilot_config(
        client_id, paused=body.paused, timezone=body.timezone, content_types=normalized
    )
    if normalized is not None:
        disabled = [f for f, e in normalized.items() if not e["enabled"]]
        storage.skip_queue_for_formats(client_id, disabled)
    return _autopilot_config_payload(client_id)


def _queue_item_or_404(client_id: str, queue_id: str) -> dict:
    item = storage.get_queue_item(client_id, queue_id)
    if not item:
        raise HTTPException(status_code=404, detail=f"queue item '{queue_id}' not found")
    return item


def _vetoable_or_422(item: dict) -> None:
    """Veto actions are allowed only while the item is upcoming."""
    if item["state"] not in ("pending", "approved"):
        raise HTTPException(status_code=422, detail="this topic is no longer actionable")
    if item["veto_expires_at"] <= datetime.now(dt_timezone.utc):
        raise HTTPException(status_code=422, detail="the veto window for this topic has closed")


def _queue_approve(client_id: str, queue_id: str) -> dict:
    item = _queue_item_or_404(client_id, queue_id)
    _vetoable_or_422(item)
    storage.update_queue_item(queue_id, state="approved")
    return {**item, "state": "approved"}


def _queue_skip(client_id: str, queue_id: str) -> dict:
    item = _queue_item_or_404(client_id, queue_id)
    _vetoable_or_422(item)
    storage.update_queue_item(queue_id, state="skipped")
    return {**item, "state": "skipped"}


def _queue_swap(client_id: str, queue_id: str) -> dict:
    item = _queue_item_or_404(client_id, queue_id)
    _vetoable_or_422(item)
    if int(item.get("swap_count") or 0) >= 1:
        raise HTTPException(status_code=422, detail="this topic was already swapped once — approve or skip it")
    pool = storage.unused_suggestions(client_id, item["format"], limit=1)
    if not pool:
        raise HTTPException(
            status_code=422,
            detail="no alternative topics available — generate topic ideas for this format first",
        )
    alt = pool[0]
    storage.update_queue_item(
        queue_id, topic=alt["topic"], suggestion_id=str(alt["id"]), swap_count=1
    )
    return {**item, "topic": alt["topic"], "suggestion_id": str(alt["id"]), "swap_count": 1}


@api.get("/autopilot/config")
def autopilot_config(client_id: str = Query(...)):
    return jsonable(_autopilot_config_payload(client_id))


@api.put("/autopilot/config")
def autopilot_config_update(body: AutopilotConfigUpdate, client_id: str = Query(...)):
    return jsonable(_update_autopilot_config(client_id, body))


@api.post("/autopilot/pause")
def autopilot_pause(client_id: str = Query(...)):
    return jsonable(_update_autopilot_config(client_id, AutopilotConfigUpdate(paused=True)))


@api.post("/autopilot/resume")
def autopilot_resume(client_id: str = Query(...)):
    return jsonable(_update_autopilot_config(client_id, AutopilotConfigUpdate(paused=False)))


@api.get("/autopilot/queue")
def autopilot_queue(client_id: str = Query(...)):
    return jsonable(storage.list_autopilot_queue(client_id))


@api.post("/autopilot/queue/{queue_id}/approve")
def autopilot_queue_approve(queue_id: str, client_id: str = Query(...)):
    return jsonable(_queue_approve(client_id, queue_id))


@api.post("/autopilot/queue/{queue_id}/swap")
def autopilot_queue_swap(queue_id: str, client_id: str = Query(...)):
    return jsonable(_queue_swap(client_id, queue_id))


@api.post("/autopilot/queue/{queue_id}/skip")
def autopilot_queue_skip(queue_id: str, client_id: str = Query(...)):
    return jsonable(_queue_skip(client_id, queue_id))


@api.get("/agent-logs")
async def agent_logs():
    async def event_stream():
        if not os.path.exists(LOG_PATH):
            open(LOG_PATH, "w").close()
        with open(LOG_PATH, "r", encoding="utf-8", errors="replace") as f:
            while True:
                line = f.readline()
                if line:
                    yield f"data: {json.dumps({'text': line})}\n\n"
                else:
                    proc = _run_state["process"]
                    if proc is None or proc.poll() is not None:
                        # read any remaining lines before closing
                        line = f.readline()
                        if line:
                            yield f"data: {json.dumps({'text': line})}\n\n"
                            continue
                        yield "event: close\ndata: {}\n\n"
                        break
                    await asyncio.sleep(0.5)
                    # Clear internal EOF buffer state
                    f.seek(f.tell())

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


def jsonable(obj):
    """Encode datetimes etc. for JSONResponse bodies built by hand."""
    from fastapi.encoders import jsonable_encoder

    return jsonable_encoder(obj)


# --------------------------------------------------------------------------- #
# App
# --------------------------------------------------------------------------- #
@asynccontextmanager
async def lifespan(app: FastAPI):
    # Ensure the schema exists before serving traffic.
    init_schema()
    # Seed the content-type/format catalog from code defaults if empty.
    try:
        storage.seed_registry_defaults()
    except Exception as e:
        print(f"[startup] WARNING: could not seed content-type/format catalog: {e}")
    # Seed the default client's profile v1 if it has none yet (idempotent;
    # profile edits afterwards live only in the DB as append-only versions).
    try:
        from casinogurus_ai_content_engine___daily_5_topic_batch.profile import load_seed_client

        record = load_seed_client("casinogurus")
        storage.upsert_client(record.client_id, record.display_name, record.site_domain)
        if storage.get_client(record.client_id)["profile_version"] == 0:
            storage.insert_profile_version(record.client_id, record.profile.model_dump(), created_by="lifespan-seed")
    except Exception as e:  # seeding must never block serving
        print(f"[startup] WARNING: could not seed default client profile: {e}")

    # Autopilot scheduler: an in-process loop ticking ~60s. All state is in
    # Postgres so restarts resume cleanly; the tick only launches work when
    # the engine is idle, so manual runs always come first. Kill-switch:
    # AUTOPILOT_DISABLED=1 (local dev / tests).
    ap_task = None
    if os.environ.get("AUTOPILOT_DISABLED", "").strip().lower() not in ("1", "true", "yes"):

        def _engine_idle() -> bool:
            proc = _run_state["process"]
            return proc is None or proc.poll() is not None

        def _launch_suggest(client_id: str, content_type: str, fmt: str) -> dict:
            return _start_agent_run(client_id, content_type, fmt, None, kind="suggest", origin="autopilot")

        def _launch_generate(item: dict) -> dict:
            return _start_agent_run(
                item["client_id"], item["content_type"], item["format"], None,
                kind="generate", topics=[item["topic"]], origin="autopilot",
            )

        async def _autopilot_loop():
            while True:
                try:
                    action = await asyncio.to_thread(
                        autopilot.tick, _engine_idle, _launch_suggest, _launch_generate
                    )
                    if action not in ("idle", "busy"):
                        print(f"[autopilot] {action}", flush=True)
                except Exception as e:
                    print(f"[autopilot] tick error: {e}", flush=True)
                await asyncio.sleep(60)

        ap_task = asyncio.create_task(_autopilot_loop())
        print("[startup] autopilot scheduler running (AUTOPILOT_DISABLED=1 to turn off)")

    yield

    if ap_task:
        ap_task.cancel()


def _cors_origins() -> list[str]:
    raw = os.environ.get("FRONTEND_ORIGIN", "*")
    origins = [o.strip() for o in raw.split(",") if o.strip()]
    if not origins:
        origins = ["*"]
    
    # Always allow these origins just in case environment variables aren't updated
    safe_defaults = [
        "https://content-agent-bice.vercel.app",
        "http://localhost:3000",
    ]
    for d in safe_defaults:
        if d not in origins:
            origins.append(d)
            
    return origins


app = FastAPI(title="NEXUS Content Engine API", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=_cors_origins(),
    allow_credentials=False,  # token-based auth; no cookies
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/healthz")
def healthz():
    return {"status": "ok"}


app.include_router(api)
app.include_router(portal)
