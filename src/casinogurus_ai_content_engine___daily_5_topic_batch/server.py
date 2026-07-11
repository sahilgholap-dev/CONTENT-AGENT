"""Tiny local HTTP server that exposes the SQLite content store to the viewer.

A browser cannot read a SQLite file directly, so this stdlib-only server sits in
front of `content_engine.db` and serves both the viewer HTML and a small JSON API
that `package-viewer.html` fetches from.

Run it, then open http://localhost:8000/ :

    # from the src/ directory
    python -m casinogurus_ai_content_engine___daily_5_topic_batch.server
    # optional: custom port / db
    python -m casinogurus_ai_content_engine___daily_5_topic_batch.server --port 8080 --db ../content_engine.db

Endpoints
---------
    GET /                     -> serves package-viewer.html
    GET /api/batches          -> [{id, batch_date, counts, source, ingested_at, package_count}, ...]
    GET /api/batches/<id>     -> the full batch object (exactly the shape the viewer renders)
    GET /api/latest           -> the most recently ingested batch object
    GET  /api/packages/<pid>/image        -> stored featured image (200 with b64, else 404)
    POST /api/packages/<pid>/image[?force=1] -> generate + store the image, return it
"""

from __future__ import annotations

import argparse
import json
import os
import time
import subprocess
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer

from casinogurus_ai_content_engine___daily_5_topic_batch.storage import (
    DEFAULT_DB_PATH,
    connect,
    get_image,
    _PROJECT_ROOT,
)

HTML_PATH = os.path.join(_PROJECT_ROOT, "package-viewer.html")
LOG_PATH = os.path.join(_PROJECT_ROOT, "agent.log")
global_agent_process = None


def _list_batches(db_path: str) -> list[dict]:
    conn = connect(db_path)
    try:
        rows = conn.execute(
            """SELECT b.id, b.batch_date, b.total_packages, b.ready_for_review_count,
                      b.needs_review_count, b.source, b.ingested_at,
                      COUNT(p.package_id) AS package_count
               FROM batches b
               LEFT JOIN packages p ON p.batch_id = b.id
               GROUP BY b.id
               ORDER BY b.ingested_at DESC, b.id DESC"""
        ).fetchall()
        return [dict(r) for r in rows]
    finally:
        conn.close()


def _get_batch(db_path: str, batch_id: int) -> dict | None:
    conn = connect(db_path)
    try:
        row = conn.execute(
            "SELECT id, raw_json FROM batches WHERE id = ?", (batch_id,)
        ).fetchone()
        if not row:
            return None
        batch = json.loads(row["raw_json"])
        if isinstance(batch, dict):
            batch["id"] = row["id"]
        return batch
    finally:
        conn.close()


def _latest_batch(db_path: str) -> dict | None:
    conn = connect(db_path)
    try:
        row = conn.execute(
            "SELECT id, raw_json FROM batches ORDER BY ingested_at DESC, id DESC LIMIT 1"
        ).fetchone()
        if not row:
            return None
        batch = json.loads(row["raw_json"])
        if isinstance(batch, dict):
            batch["id"] = row["id"]
        return batch
    finally:
        conn.close()


def _image_pid(path: str) -> str | None:
    """Extract <pid> from /api/packages/<pid>/image, else None."""
    prefix, suffix = "/api/packages/", "/image"
    if path.startswith(prefix) and path.endswith(suffix):
        pid = path[len(prefix) : -len(suffix)]
        return pid or None
    return None


def _image_payload(row: dict | None) -> dict | None:
    if not row:
        return None
    return {
        "package_id": row.get("package_id"),
        "status": row.get("status"),
        "image_b64": row.get("image_b64"),
        "mime_type": row.get("mime_type"),
        "alt_text": row.get("alt_text"),
        "prompt": row.get("prompt"),
        "model": row.get("model"),
        "size": row.get("size"),
        "error": row.get("error"),
        "created_at": row.get("created_at"),
    }


def make_handler(db_path: str):
    class Handler(BaseHTTPRequestHandler):
        # Quieter logging
        def log_message(self, fmt, *args):
            pass

        def _send_json(self, obj, status=200):
            body = json.dumps(obj, ensure_ascii=False).encode("utf-8")
            self.send_response(status)
            self.send_header("Content-Type", "application/json; charset=utf-8")
            self.send_header("Access-Control-Allow-Origin", "*")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)

        def _send_html(self):
            try:
                with open(HTML_PATH, "rb") as fh:
                    body = fh.read()
            except FileNotFoundError:
                self._send_json({"error": f"HTML not found at {HTML_PATH}"}, 500)
                return
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)

        def _send_zip(self, batch: dict):
            import io
            import zipfile
            from bs4 import BeautifulSoup
            import docx

            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
                for idx, pkg in enumerate(batch.get("packages", [])):
                    doc = docx.Document()
                    title = pkg.get("topic") or pkg.get("primary_keyword") or f"Topic_{idx+1}"
                    doc.add_heading(title, 0)

                    doc.add_heading("Metadata", level=1)
                    doc.add_paragraph(f"Primary Keyword: {pkg.get('primary_keyword', '')}")
                    
                    draft = pkg.get("draft", {})
                    doc.add_paragraph(f"Meta Description: {draft.get('meta_description', '')}")
                    
                    doc.add_heading("Content", level=1)
                    html_content = draft.get("body_html", "")
                    if html_content:
                        soup = BeautifulSoup(html_content, "html.parser")
                        for element in soup.find_all(['h2', 'h3', 'p', 'ul', 'ol', 'li']):
                            if element.name == 'h2':
                                doc.add_heading(element.get_text(), level=2)
                            elif element.name == 'h3':
                                doc.add_heading(element.get_text(), level=3)
                            elif element.name == 'p':
                                doc.add_paragraph(element.get_text())
                            elif element.name in ['ul', 'ol']:
                                pass
                            elif element.name == 'li':
                                doc.add_paragraph(element.get_text(), style='List Bullet')

                    flags = draft.get("verification_flags", [])
                    if flags:
                        doc.add_heading("Verification Flags", level=1)
                        for f in flags:
                            flag_str = f.get('flag') if isinstance(f, dict) else str(f)
                            doc.add_paragraph(flag_str, style='List Bullet')

                    compliance = pkg.get("compliance_scorecard", {})
                    if compliance:
                        doc.add_heading("Compliance Checks", level=1)
                        doc.add_paragraph(f"Overall Verdict: {compliance.get('overall_verdict', 'Unknown')}")
                        
                        blocking = compliance.get("blocking_failures", [])
                        if blocking:
                            doc.add_heading("Blocking Failures", level=2)
                            for b in blocking:
                                if isinstance(b, str):
                                    doc.add_paragraph(b, style='List Bullet')
                                else:
                                    name = b.get('check_name') or b.get('item') or b.get('check') or b.get('name') or "Unknown Check"
                                    sev = b.get('severity')
                                    sev_str = f" ({sev})" if sev else ""
                                    doc.add_heading(f"{name}{sev_str}", level=3)
                                    if b.get('violation'):
                                        doc.add_paragraph(f"Violation: {b.get('violation')}")
                                    if b.get('remediation'):
                                        doc.add_paragraph(f"Remediation: {b.get('remediation')}")
                                        
                        checks = compliance.get("checks", [])
                        if checks:
                            doc.add_heading("All Checks", level=2)
                            for c in checks:
                                name = c.get('check_name') or c.get('item') or c.get('check') or c.get('name') or "Unknown Check"
                                verdict = c.get('verdict') or c.get('result') or ""
                                sev = c.get('severity')
                                sev_str = f" [{sev}]" if sev else ""
                                doc.add_paragraph(f"{name}{sev_str}: {verdict}", style='List Bullet')
                                details = c.get('offending_text') or c.get('offendingText') or c.get('violation') or c.get('details') or ""
                                if details:
                                    # We don't have docx.shared imported, just add a regular paragraph
                                    doc.add_paragraph(f"   Details: {details}")

                    doc_buffer = io.BytesIO()
                    doc.save(doc_buffer)
                    
                    # Clean filename
                    filename = "".join([c for c in title if c.isalpha() or c.isdigit() or c==' ']).rstrip()
                    filename = filename.replace(" ", "_") + ".docx"
                    zip_file.writestr(filename, doc_buffer.getvalue())

            zip_buffer.seek(0)
            body = zip_buffer.read()
            self.send_response(200)
            self.send_header("Content-Type", "application/zip")
            batch_date = batch.get("batch_date", "download")
            self.send_header("Content-Disposition", f'attachment; filename="batch_{batch_date}.zip"')
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)

        def do_GET(self):
            path = self.path.split("?", 1)[0].rstrip("/")

            if path in ("", "/index.html", "/package-viewer.html"):
                return self._send_html()

            if path == "/api/batches":
                return self._send_json(_list_batches(db_path))

            if path == "/api/latest":
                batch = _latest_batch(db_path)
                if batch is None:
                    return self._send_json({"error": "no batches stored"}, 404)
                return self._send_json(batch)

            if path == "/api/agent-logs":
                self.send_response(200)
                self.send_header("Content-Type", "text/event-stream")
                self.send_header("Cache-Control", "no-cache")
                self.send_header("Connection", "keep-alive")
                self.send_header("Access-Control-Allow-Origin", "*")
                self.send_header("X-Accel-Buffering", "no")
                self.end_headers()

                if not os.path.exists(LOG_PATH):
                    open(LOG_PATH, "w").close()

                global global_agent_process
                with open(LOG_PATH, "r", encoding="utf-8", errors="replace") as f:
                    while True:
                        line = f.readline()
                        if line:
                            data = json.dumps({"text": line})
                            try:
                                self.wfile.write(f"data: {data}\n\n".encode("utf-8"))
                                self.wfile.flush()
                            except (BrokenPipeError, ConnectionResetError):
                                break
                        else:
                            if global_agent_process is None or global_agent_process.poll() is not None:
                                try:
                                    self.wfile.write(b"event: close\ndata: {}\n\n")
                                    self.wfile.flush()
                                except (BrokenPipeError, ConnectionResetError):
                                    pass
                                break
                            time.sleep(0.5)
                return

            pid = _image_pid(path)
            if pid is not None:
                row = get_image(pid, db_path=db_path)
                payload = _image_payload(row)
                if payload is None or payload.get("status") != "ok" or not payload.get("image_b64"):
                    # 404 with any error detail so the UI can show "generate" or the failure.
                    body = payload or {"package_id": pid, "status": "none"}
                    return self._send_json(body, 404)
                return self._send_json(payload)

            if path.startswith("/api/batches/"):
                if path.endswith("/download"):
                    raw = path.split("/")[-2]
                    try:
                        bid = int(raw)
                    except ValueError:
                        return self._send_json({"error": "batch id must be an integer"}, 400)
                    batch = _get_batch(db_path, bid)
                    if batch is None:
                        return self._send_json({"error": f"batch {bid} not found"}, 404)
                    return self._send_zip(batch)

                raw = path.rsplit("/", 1)[-1]
                try:
                    bid = int(raw)
                except ValueError:
                    return self._send_json({"error": "batch id must be an integer"}, 400)
                batch = _get_batch(db_path, bid)
                if batch is None:
                    return self._send_json({"error": f"batch {bid} not found"}, 404)
                return self._send_json(batch)

            self._send_json({"error": "not found", "path": path}, 404)

        def do_POST(self):
            path = self.path.split("?", 1)[0].rstrip("/")
            query = self.path.split("?", 1)[1] if "?" in self.path else ""
            force = "force=1" in query or "force=true" in query

            pid = _image_pid(path)
            if pid is None:
                if path == "/api/run-agent":
                    import sys
                    global global_agent_process
                    if global_agent_process is not None and global_agent_process.poll() is None:
                        return self._send_json({"error": "Agent is already running"}, 400)
                    
                    # Kick off the crew asynchronously
                    try:
                        log_file = open(LOG_PATH, "w", encoding="utf-8")
                        cmd = [sys.executable, "-u", "-m", "casinogurus_ai_content_engine___daily_5_topic_batch.main", "run"]
                        global_agent_process = subprocess.Popen(
                            cmd, 
                            cwd=_PROJECT_ROOT, 
                            stdout=log_file, 
                            stderr=subprocess.STDOUT
                        )
                        return self._send_json({"status": "started", "message": "Agent execution started in background."})
                    except Exception as e:
                        return self._send_json({"error": str(e)}, 500)
                return self._send_json({"error": "not found", "path": path}, 404)

            # Lazy import so a missing openai package only bites when generating.
            try:
                from casinogurus_ai_content_engine___daily_5_topic_batch.images import (
                    generate_for_package,
                )
                row = generate_for_package(pid, force=force, db_path=db_path)
            except ValueError as e:  # unknown package id
                return self._send_json({"error": str(e)}, 404)
            except Exception as e:  # unexpected
                return self._send_json({"error": str(e)}, 500)

            payload = _image_payload(row)
            status = 200 if payload and payload.get("status") == "ok" else 502
            return self._send_json(payload or {"package_id": pid, "status": "error"}, status)

    return Handler


def serve(db_path: str = DEFAULT_DB_PATH, host: str = "127.0.0.1", port: int = 8000) -> None:
    handler = make_handler(db_path)
    httpd = ThreadingHTTPServer((host, port), handler)
    print(f"Content viewer serving on http://{host}:{port}/")
    print(f"  DB:   {db_path}")
    print(f"  HTML: {HTML_PATH}")
    print("Press Ctrl+C to stop.")
    try:
        httpd.serve_forever()
    except KeyboardInterrupt:
        print("\nStopping server.")
    finally:
        httpd.server_close()


def _main() -> None:
    ap = argparse.ArgumentParser(description="Serve the content store to the viewer.")
    ap.add_argument("--db", default=DEFAULT_DB_PATH, help="Path to the SQLite database.")
    ap.add_argument("--host", default="127.0.0.1")
    ap.add_argument("--port", type=int, default=8000)
    args = ap.parse_args()
    serve(db_path=args.db, host=args.host, port=args.port)


if __name__ == "__main__":
    _main()
